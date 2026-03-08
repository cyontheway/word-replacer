#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 敏感词替换工具 - FastAPI 后端 (v3)
新增：预览对比 + 预检报告功能
"""

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from docx import Document
import pandas as pd
import re
import io
import os
import tempfile
import uuid
import json
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass, asdict

app = FastAPI(title="Word 敏感词替换工具", version="3.0")

# 存储临时文件与原文件名的映射
temp_file_mapping: Dict[str, str] = {}

@dataclass
class MatchDetail:
    """匹配详情，用于预览"""
    id: str  # 唯一标识符，用于前端勾选追踪
    start: int
    end: int
    original: str
    replacement: str
    match_type: str
    context: str

# 全局匹配项 ID 计数器
_match_id_counter = 0

def get_next_match_id() -> str:
    """生成下一个匹配项 ID"""
    global _match_id_counter
    _match_id_counter += 1
    return f"match_{_match_id_counter}"

def reset_match_id_counter():
    """重置匹配项 ID 计数器（每次预览前调用）"""
    global _match_id_counter
    _match_id_counter = 0

def get_context(text: str, start: int, end: int, context_size: int = 15) -> str:
    """获取匹配项的上下文"""
    context_start = max(0, start - context_size)
    context_end = min(len(text), end + context_size)
    return text[context_start:context_end]

# ============ 内置规则配置 ============
BANK_KEYWORDS = [
    "中国工商银行", "工商银行", "工行",
    "中国农业银行", "农业银行", "农行",
    "中国银行", "中行",
    "中国建设银行", "建设银行", "建行",
    "交通银行", "交行",
    "中国邮政储蓄银行", "邮储银行",
    "招商银行", "招行",
    "上海浦东发展银行", "浦发银行", "浦发",
    "中国民生银行", "民生银行",
    "中信银行",
    "兴业银行",
    "中国光大银行", "光大银行",
    "平安银行",
    "广发银行",
    "华夏银行",
    "浙商银行",
]

COMPANY_SUFFIXES = r'(?:有限公司|有限责任公司|股份公司|股份有限公司|集团(?:公司)?|企业|事务所|中心)'
EXCLUDE_PREFIXES = ['根据', '鉴于', '按照', '依据', '如果', '由于']

# 商号/品牌名（在合同中可能作为甲乙方代称）
BRAND_NAMES = [
    "淘宝", "天猫", "阿里巴巴", "支付宝", "菜鸟", "盒马",
    "星巴克", "麦当劳", "肯德基", "必胜客",
    "京东", "拼多多", "美团", "饿了么", "滴滴", "高德",
    "腾讯", "微信", "QQ", "王者荣耀",
    "百度", "字节", "抖音", "快手", "B站", "哔哩哔哩",
    "网易", "新浪", "微博", "知乎", "小红书",
    "苹果", "华为", "小米", "OPPO", "vivo", "三星",
    "索尼", "微软", "谷歌", "亚马逊",
    "海尔", "美的", "格力", "TCL", "海信",
    "顺丰", "圆通", "中通", "申通", "韵达", "EMS",
    "中国移动", "中国联通", "中国电信",
]

# 需要排除的高频合同词汇（不视为公司名）
CONTRACT_COMMON_WORDS = [
    '甲方', '乙方', '卖方', '买方', '供方', '需方', '出租方', '承租方',
    '发包方', '承包方', '供应商', '客户', '出卖人', '买受人', '委托人',
    '受托人', '承揽人', '定作人', '债权人', '债务人', '保证人', '被保证人',
    '权利人', '义务人', '所有人', '使用人', '受益人', '投保人', '被保险人',
    '第三人', '相对人', '行为人', '当事人', '见证人', '代理人', '被代理人',
    '知识产权', '所有权', '使用权', '经营权', '管理权', '决策权',
    '违约金', '赔偿金', '定金', '预付款', '尾款', '价款', '报酬', '费用',
    '交付', '验收', '安装', '调试', '培训', '维护', '保修', '质保',
    '保密', '保密义务', '竞业限制', '知识产权', '商业秘密',
]

# ============ 核心替换函数（返回详细信息） ============
def replace_bank_names(text: str) -> tuple:
    """替换银行名称，返回 (新文本, 计数, 匹配详情列表)"""
    count = 0
    matches = []
    offset = 0

    for bank in BANK_KEYWORDS:
        pattern = re.compile(re.escape(bank), re.IGNORECASE)
        for m in pattern.finditer(text):
            original = m.group()
            start = m.start() + offset
            end = m.end() + offset

            # 检查原文是否已经被【】包裹
            context_before = text[max(0, start-1):start]
            context_after = text[end:min(len(text), end+1)]
            if context_before == '【' and context_after == '】':
                replacement = "XX银行"  # 不加外层【】
            else:
                replacement = "【XX银行】"

            matches.append(MatchDetail(
                id=get_next_match_id(),
                start=start,
                end=end,
                original=original,
                replacement=replacement,
                match_type="银行名",
                context=get_context(text, start, end)
            ))

            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)
            count += 1

    return text, count, matches

def replace_urls(text: str) -> tuple:
    """替换URL/网址，返回 (新文本, 计数, 匹配详情列表)"""
    count = 0
    matches = []
    offset = 0

    # URL正则：匹配 http:// 或 https:// 开头的URL
    url_pattern = re.compile(
        r'https?://[^\s\u3000\uff0c\u3002\uff1b\uff0c\uff1a\"\'\'\'\(\)\uff08\uff09\u300a\u300b<>\{\}\[\]]+',
        re.IGNORECASE
    )

    for m in url_pattern.finditer(text):
        original = m.group()
        start = m.start() + offset
        end = m.end() + offset
        replacement = "【URL】"

        matches.append(MatchDetail(
            id=get_next_match_id(),
            start=start,
            end=end,
            original=original,
            replacement=replacement,
            match_type="网址",
            context=get_context(text, start, end)
        ))

        text = text[:start] + replacement + text[end:]
        offset += len(replacement) - len(original)
        count += 1

    return text, count, matches

def replace_brand_names(text: str) -> tuple:
    """替换商号/品牌名（排除法+出现次数≥10次），返回 (新文本, 计数, 匹配详情列表)"""
    count = 0
    matches = []
    offset = 0

    # 排除的合同常见词
    EXCLUDE_WORDS = {
        '甲方', '乙方', '丙方', '丁方', '买方', '卖方', '供方', '需方',
        '承揽人', '服务方', '供应商', '买方', '出租方', '承租方',
        '发包方', '承包方', '委托方', '受托方', '出卖人', '买受人',
        '权利', '义务', '责任', '违约', '合同', '协议', '条款', '约定',
        '订单', '解除', '终止'
    }

    # 找出所有2-15字符的中文词组
    word_pattern = re.compile(r'[\u4e00-\u9fa5]{2,15}')
    word_counts = {}

    for m in word_pattern.finditer(text):
        word = m.group()
        if word not in EXCLUDE_WORDS:
            word_counts[word] = word_counts.get(word, 0) + 1

    # 筛选出现次数≥10次的词
    frequent_words = {word: cnt for word, cnt in word_counts.items() if cnt >= 10}

    # 替换这些高频词
    for brand, freq in frequent_words.items():
        pattern = re.compile(re.escape(brand), re.IGNORECASE)
        for m in pattern.finditer(text):
            original = m.group()
            start = m.start() + offset
            end = m.end() + offset

            # 检查是否被【】包裹
            context_before = text[max(0, start-1):start]
            context_after = text[end:min(len(text), end+1)]

            # 检查前后文是否是合同高频词汇（如果是则跳过）
            # 向后看最多4个字符，看是否组成合同常用词
            context_after_more = text[end:min(len(text), end+4)]
            if any(word.startswith(original) or original.startswith(word[:2]) for word in CONTRACT_COMMON_WORDS if len(word) >= 2):
                # 进一步检查：如果匹配的是合同高频词的一部分，跳过
                matched_as_common = False
                for common_word in CONTRACT_COMMON_WORDS:
                    if common_word in original or original in common_word:
                        matched_as_common = True
                        break
                if matched_as_common:
                    continue

            if context_before == '【' and context_after == '】':
                replacement = "XX公司"
            else:
                replacement = "【XX公司】"

            matches.append(MatchDetail(
                id=get_next_match_id(),
                start=start,
                end=end,
                original=original,
                replacement=replacement,
                match_type="商号",
                context=get_context(text, start, end)
            ))

            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)
            count += 1

    return text, count, matches

def replace_company_names(text: str) -> tuple:
    """替换公司名称，返回 (新文本, 计数, 匹配详情列表)"""
    COMPANY_NAME_BODY = r'[\u4e00-\u9fa5a-zA-Z0-9()（）]{2,20}'

    pattern1 = re.compile(
        r'(甲方|乙方|卖方|买方|供方|需方|出租方|承租方|发包方|承包方|供应商|客户|出卖人|买受人)' +
        r'[：:\s]*' +
        r'(' + COMPANY_NAME_BODY + COMPANY_SUFFIXES + r')',
        re.IGNORECASE
    )

    pattern2 = re.compile(
        r'[""''《（(【](' + COMPANY_NAME_BODY + r')' + COMPANY_SUFFIXES + r'[)）》''""】]',
        re.IGNORECASE
    )

    pattern3 = re.compile(
        r'(?:^|[\s,，.。;；:：!！?？\n])' +
        r'(' + COMPANY_NAME_BODY + r')' + COMPANY_SUFFIXES +
        r'(?![\u4e00-\u9fa5a-zA-Z])',
        re.IGNORECASE
    )

    count = 0
    matches = []
    offset = 0

    for match in pattern1.finditer(text):
        company_name = match.group(2)
        prefix = match.group(1)
        if not any(company_name.startswith(exclude) for exclude in EXCLUDE_PREFIXES):
            original = match.group(0)
            start = match.start() + offset
            end = match.end() + offset
            replacement = prefix + "【XXXXXXXX】"

            matches.append(MatchDetail(
                id=get_next_match_id(),
                start=start,
                end=end,
                original=original,
                replacement=replacement,
                match_type="公司名",
                context=get_context(text, start, end)
            ))

            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)
            count += 1

    for match in pattern2.finditer(text):
        original = match.group(0)
        start = match.start() + offset
        end = match.end() + offset
        replacement = '"【XXXXXXXX】"'

        matches.append(MatchDetail(
            id=get_next_match_id(),
            start=start,
            end=end,
            original=original,
            replacement=replacement,
            match_type="公司名",
            context=get_context(text, start, end)
        ))

        text = text[:start] + replacement + text[end:]
        offset += len(replacement) - len(original)
        count += 1

    for match in pattern3.finditer(text):
        company_name = match.group(1)
        if not any(company_name.startswith(exclude) for exclude in EXCLUDE_PREFIXES):
            original = match.group(0)
            start = match.start() + offset
            end = match.end() + offset

            prefix_char = original[0]
            if prefix_char in ' \t\n,，.。;；:：!！?？':
                replacement = prefix_char + "【XXXXXXXX】"
            else:
                replacement = "【XXXXXXXX】"

            matches.append(MatchDetail(
                id=get_next_match_id(),
                start=start,
                end=end,
                original=original,
                replacement=replacement,
                match_type="公司名",
                context=get_context(text, start, end)
            ))

            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)
            count += 1

    return text, count, matches

def replace_addresses(text: str) -> tuple:
    """替换地址信息，返回 (新文本, 计数, 匹配详情列表)"""
    count = 0
    matches = []
    offset = 0

    province_names = r'(?:北京|天津|上海|重庆|河北|山西|辽宁|吉林|黑龙江|江苏|浙江|安徽|福建|江西|山东|河南|湖北|湖南|广东|海南|四川|贵州|云南|陕西|甘肃|青海|台湾|内蒙古|广西|西藏|宁夏|新疆|香港|澳门|.*?省|.*?自治区)'
    city_pattern = r'[^省市区县街道镇乡]{1,20}(?:市|州|盟)'
    district_pattern = r'[^省市区县街道镇乡]{1,20}(?:区|县|旗)'
    street_pattern = r'[^省市区县街道镇乡路街巷号栋幢单元层楼房室房座]{1,20}(?:街道|镇|乡)'
    road_pattern = r'[^省市区县街道镇乡路街巷号栋幢单元层楼房室房座]{1,20}(?:路|街|道|巷|弄|胡同)'

    trigger_pattern = re.compile(
        r'(地址|住所|注册地址?|经营场所|送达地址|联系地址|收货地址|邮寄地址)(位于)?[：:\s]*' +
        r'(' + province_names + r')?' +
        r'(' + city_pattern + r')?' +
        r'(' + district_pattern + r')?' +
        r'(' + street_pattern + r')?' +
        r'(' + road_pattern + r')' +
        r'([^，。；,;\n（(【\]]{0,30}(?:号|栋|幢|单元|层|楼|室|房|座|\d+))?' +
        r'(?=[，。；,;\n（(【]|$)',
        re.UNICODE
    )

    pure_address_pattern = re.compile(
        r'(' + province_names + r')' +
        r'(' + city_pattern + r')' +
        r'(' + district_pattern + r')' +
        r'([^，。；,;\n]{0,50}(?:街道|镇|乡))?' +
        r'([^，。；,;\n]{0,50}(?:路|街|道|巷|弄|胡同))?' +
        r'([^，。；,;\n]{0,30}(?:号|栋|幢|单元|层|楼|室|房|座|\d+))?' +
        r'(?=[，。；,;\n]|$)',
        re.UNICODE
    )

    # 方括号内地址：地址：[省市区路号]
    bracket_address_pattern = re.compile(
        r'(地址|住所|注册地址?|经营场所|送达地址|联系地址|收货地址|邮寄地址)(位于)?[：:\s]*\[' +
        r'(' + province_names + r')?' +
        r'(' + city_pattern + r')?' +
        r'(' + district_pattern + r')?' +
        r'(' + street_pattern + r')?' +
        r'(' + road_pattern + r')' +
        r'([^\]]{0,30}(?:号|栋|幢|单元|层|楼|室|房|座|\d+))?' +
        r'\]',
        re.UNICODE
    )

    def build_replacement(match_groups, has_trigger=False):
        parts = []
        # has_trigger=True时，groups是：(触发词, 位于?, 省, 市, 区, 街道, 路, 号)
        # has_trigger=False时，groups是：(省, 市, 区, 街道, 路, 号)
        offset = 2 if has_trigger else 0

        if match_groups[0 + offset]:
            parts.append("xx省")
        if match_groups[1 + offset]:
            parts.append("xx市")
        if match_groups[2 + offset]:
            parts.append("xx区")
        if match_groups[3 + offset]:
            parts.append("xx街道")
        if match_groups[4 + offset]:
            parts.append("xx路")
        if match_groups[5 + offset]:
            parts.append("xx号")

        return "".join(parts) if parts else "【XX地址】"

    for match in list(trigger_pattern.finditer(text)):
        original = match.group(0)
        start = match.start() + offset
        end = match.end() + offset

        replacement = build_replacement(match.groups(), has_trigger=True)
        trigger_word = match.group(1)
        weiyuword = match.group(2) if match.group(2) else ""
        new_str = f"{trigger_word}{weiyuword}：{replacement}"

        matches.append(MatchDetail(
            id=get_next_match_id(),
            start=start,
            end=end,
            original=original,
            replacement=new_str,
            match_type="地址",
            context=get_context(text, start, end)
        ))

        text = text[:start] + new_str + text[end:]
        offset += len(new_str) - len(original)
        count += 1

    for match in list(pure_address_pattern.finditer(text)):
        original = match.group(0)
        if '【' in original or 'xx' in original:
            continue

        start = match.start() + offset
        end = match.end() + offset
        replacement = build_replacement(match.groups(), has_trigger=False)

        matches.append(MatchDetail(
            id=get_next_match_id(),
            start=start,
            end=end,
            original=original,
            replacement=replacement,
            match_type="地址",
            context=get_context(text, start, end)
        ))

        text = text[:start] + replacement + text[end:]
        offset += len(replacement) - len(original)
        count += 1

    # 处理方括号内的地址：地址：[省市区路号]
    for match in list(bracket_address_pattern.finditer(text)):
        original = match.group(0)
        if '【' in original or 'xx' in original:
            continue

        start = match.start() + offset
        end = match.end() + offset
        replacement_addr = build_replacement(match.groups(), has_trigger=True)
        trigger_word = match.group(1)
        weiyuword = match.group(2) if match.group(2) else ""
        new_str = f"{trigger_word}{weiyuword}：[{replacement_addr}]"

        matches.append(MatchDetail(
            id=get_next_match_id(),
            start=start,
            end=end,
            original=original,
            replacement=new_str,
            match_type="地址",
            context=get_context(text, start, end)
        ))

        text = text[:start] + new_str + text[end:]
        offset += len(new_str) - len(original)
        count += 1

    return text, count, matches

def apply_regex_replacements(text: str, return_details: bool = False) -> tuple:
    """应用正则表达式替换

    Args:
        text: 待处理文本
        return_details: 是否返回匹配详情列表（用于预览）

    Returns:
        如果 return_details=False: (替换后文本, 总数, 统计字典)
        如果 return_details=True: (替换后文本, 总数, 统计字典, 匹配详情列表)
    """
    original_text = text  # 保存原文用于提取上下文
    total_count = 0
    stats = {}
    all_matches = []
    match_details = []  # 存储 MatchDetail 对象
    offset = 0  # 位置偏移量

    # 长字符串（字母+数字混合，含空格）- 超过10位的连续字母数字全部用X替换
    # 先找出所有匹配，优先保留最长的匹配，避免嵌套
    long_string_matches = []
    for m in re.finditer(r'[A-Za-z0-9][\sA-Za-z0-9]*[A-Za-z0-9]|[A-Za-z0-9]{10,}', text):
        original = m.group()
        chars_only = re.sub(r'\s', '', original)
        if len(chars_only) >= 10:
            start = m.start()
            end = m.end()
            # 检查是否已经在【】内
            if start > 0 and text[start-1] == '【':
                continue
            long_string_matches.append((start, end, original, chars_only))

    # 按长度降序排序，优先处理最长的匹配
    long_string_matches.sort(key=lambda x: -(x[1] - x[0]))

    # 去重：保留最长的匹配，丢弃与之重叠的较短匹配
    filtered_matches = []
    for start, end, original, chars_only in long_string_matches:
        is_overlapping = False
        for prev_start, prev_end, _, _ in filtered_matches:
            if not (end <= prev_start or start >= prev_end):
                is_overlapping = True
                break
        if not is_overlapping:
            filtered_matches.append((start, end, original, chars_only))

    # 按位置从后往前排序，避免替换时位置偏移
    filtered_matches.sort(key=lambda x: x[0], reverse=True)
    for start, end, original, chars_only in filtered_matches:
        replacement = "【" + "X" * len(chars_only) + "】"
        all_matches.append((start, end, replacement, "长字符串", original))
        text = text[:start] + replacement + text[end:]

    # 身份证号 - 全部用X替换
    for m in re.finditer(r'\d{17}[\dXx]|\d{15}', text):
        start = m.start() + offset
        end = m.end() + offset
        original = m.group()
        # 检查是否已经在【】内，如果是则跳过
        if start > 0 and text[start-1] == '【':
            continue
        replacement = "【XXXXXXXXXXXXXXXXXX】"

        all_matches.append((start, end, replacement, "身份证号", original))
        text = text[:start] + replacement + text[end:]
        offset += len(replacement) - len(original)

    # 统一信用代码 - 全部用X替换
    for m in re.finditer(r'[0-9A-Z]{18}', text):
        start = m.start() + offset
        end = m.end() + offset
        original = m.group()
        # 检查是否已经在【】内，如果是则跳过
        if start > 0 and text[start-1] == '【':
            continue
        if original[:2] in ['11', '12', '13', '19', '51', '52', '53', '59', '91', '92', '93', '99'] or \
           any(c.isalpha() for c in original[-10:]):
            replacement = "【XXXXXXXXXXXXXXXXXX】"
            all_matches.append((start, end, replacement, "统一社会信用代码", original))
            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)

    # 银行卡号 - 全部用X替换
    for m in re.finditer(r'\d{16,19}|(?:\d{4}[-\s]){3,4}\d{1,4}', text):
        start = m.start() + offset
        end = m.end() + offset
        original = m.group()
        # 检查是否已经在【】内，如果是则跳过
        if start > 0 and text[start-1] == '【':
            continue
        clean_card = re.sub(r'[-\s]', '', original)
        if len(clean_card) >= 16 and len(clean_card) <= 19:
            replacement = "【XXXXXXXXXXXXXXXX】"
            all_matches.append((start, end, replacement, "银行卡号", original))
            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)

    # 手机号 - 全部用X替换
    for m in re.finditer(r'1[3-9]\d{9}', text):
        start = m.start() + offset
        end = m.end() + offset
        original = m.group()
        if (start == 0 or not text[start-1].isdigit()) and \
           (end == len(text) or not text[end].isdigit()):
            replacement = "【XXXXXXXXXXX】"
            all_matches.append((start, end, replacement, "手机号", original))
            text = text[:start] + replacement + text[end:]
            offset += len(replacement) - len(original)

    # 座机电话 - 全部用X替换
    for m in re.finditer(r'0\d{2,3}-\d{7,8}|\(\d{2,4}\)\s*\d{7,8}', text):
        start = m.start() + offset
        end = m.end() + offset
        original = m.group()
        replacement = "【XXXXXXXX】"
        all_matches.append((start, end, replacement, "座机电话", original))
        text = text[:start] + replacement + text[end:]
        offset += len(replacement) - len(original)

    # 电子邮箱 - 全部用X替换
    for m in re.finditer(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text):
        start = m.start() + offset
        end = m.end() + offset
        original = m.group()
        replacement = "【***@***.com】"
        all_matches.append((start, end, replacement, "电子邮箱", original))
        text = text[:start] + replacement + text[end:]
        offset += len(replacement) - len(original)

    # 处理数字类匹配的详情
    for start, end, replacement, match_type, original in all_matches:
        if return_details:
            match_details.append(MatchDetail(
                id=get_next_match_id(),
                start=start,
                end=end,
                original=original,
                replacement=replacement,
                match_type=match_type,
                context=get_context(original_text, start, end)
            ))
        stats[match_type] = stats.get(match_type, 0) + 1
        total_count += 1

    # URL/网址
    text, url_count, url_matches = replace_urls(text)
    if url_count > 0:
        stats["网址"] = url_count
        total_count += url_count
        if return_details:
            match_details.extend(url_matches)

    # 商号/品牌名
    text, brand_count, brand_matches = replace_brand_names(text)
    if brand_count > 0:
        stats["商号"] = brand_count
        total_count += brand_count
        if return_details:
            match_details.extend(brand_matches)

    # 银行名
    text, bank_count, bank_matches = replace_bank_names(text)
    if bank_count > 0:
        stats["银行名"] = bank_count
        total_count += bank_count
        if return_details:
            match_details.extend(bank_matches)

    # 公司名
    text, company_count, company_matches = replace_company_names(text)
    if company_count > 0:
        stats["公司名"] = company_count
        total_count += company_count
        if return_details:
            match_details.extend(company_matches)

    # 地址
    text, address_count, address_matches = replace_addresses(text)
    if address_count > 0:
        stats["地址"] = address_count
        total_count += address_count
        if return_details:
            match_details.extend(address_matches)

    if return_details:
        # 按位置排序（从前到后）
        match_details.sort(key=lambda x: x.start)
        return text, total_count, stats, match_details

    return text, total_count, stats

def collect_all_matches(text: str, rules_df: pd.DataFrame = None) -> list:
    """收集所有匹配（手动规则 + 正则），返回排序后的 MatchDetail 列表"""
    all_matches = []

    # 1. 收集手动规则匹配
    if rules_df is not None and not rules_df.empty:
        for _, row in rules_df.iterrows():
            old_word = str(row['原词']) if pd.notna(row['原词']) else ""
            new_word = str(row['替换词']) if pd.notna(row['替换词']) else ""
            if old_word and old_word != 'nan':
                for m in re.finditer(re.escape(old_word), text):
                    marked_new = f"【{new_word}】" if not new_word.startswith('【') else new_word
                    all_matches.append(MatchDetail(
                        id=get_next_match_id(),
                        start=m.start(),
                        end=m.end(),
                        original=m.group(),
                        replacement=marked_new,
                        match_type="自定义规则",
                        context=get_context(text, m.start(), m.end())
                    ))

    # 2. 收集正则匹配（获取详情但不替换）
    _, _, regex_stats, regex_matches = apply_regex_replacements(text, return_details=True)
    all_matches.extend(regex_matches)

    # 3. 按位置排序，并处理重叠（保留最长匹配）
    all_matches.sort(key=lambda x: (x.start, - (x.end - x.start)))

    filtered = []
    last_end = -1
    for m in all_matches:
        if m.start >= last_end:  # 不重叠
            filtered.append(m)
            last_end = m.end

    return filtered

def process_document(doc_file, rules_df: pd.DataFrame = None, selected_match_ids: set = None) -> tuple:
    """处理 Word 文档 - 一次性替换避免套娃

    Args:
        doc_file: Word文档文件对象
        rules_df: 自定义规则DataFrame
        selected_match_ids: 用户选中的匹配项ID集合（None表示全选）
    """
    doc = Document(doc_file)
    excel_replace_count = 0
    regex_stats_total = {}

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        original_text = para.text
        if not original_text:
            continue

        # 收集所有匹配（一次性）
        matches = collect_all_matches(original_text, rules_df)

        if not matches:
            continue

        # TODO: 根据 selected_match_ids 过滤匹配项（待实现）
        # if selected_match_ids:
        #     matches = [m for m in matches if m.id in selected_match_ids]

        if not matches:
            continue

        # 统计
        for m in matches:
            if m.match_type == "自定义规则":
                excel_replace_count += 1
            else:
                regex_stats_total[m.match_type] = regex_stats_total.get(m.match_type, 0) + 1

        # 一次性替换（从后往前）
        modified_text = original_text
        for m in sorted(matches, key=lambda x: x.start, reverse=True):
            modified_text = modified_text[:m.start] + m.replacement + modified_text[m.end:]

        para.text = modified_text

    return doc, excel_replace_count, regex_stats_total

def preview_document(doc_file, rules_df: pd.DataFrame = None) -> dict:
    """预览文档替换结果（不修改文档）"""
    reset_match_id_counter()  # 重置匹配项ID计数器
    doc = Document(doc_file)
    all_matches = []
    stats = {}
    original_paragraphs = []
    replaced_paragraphs = []

    # 收集所有段落
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    global_offset = 0  # 全局偏移量，用于计算跨段落位置

    for para in all_paras:
        original_text = para.text
        if not original_text:
            continue

        original_paragraphs.append(original_text)
        modified_text = original_text
        para_matches = []

        # Excel 规则替换（收集匹配）
        if rules_df is not None and not rules_df.empty:
            for _, row in rules_df.iterrows():
                old_word = str(row['原词']) if pd.notna(row['原词']) else ""
                new_word = str(row['替换词']) if pd.notna(row['替换词']) else ""
                if old_word and old_word != 'nan':
                    for m in re.finditer(re.escape(old_word), modified_text):
                        marked_new = f"【{new_word}】"
                        para_matches.append(MatchDetail(
                            id=get_next_match_id(),
                            start=global_offset + m.start(),
                            end=global_offset + m.end(),
                            original=old_word,
                            replacement=marked_new,
                            match_type="自定义规则",
                            context=get_context(original_text, m.start(), m.end())
                        ))
                        stats["自定义规则"] = stats.get("自定义规则", 0) + 1
                    modified_text = modified_text.replace(
                        old_word,
                        f"【{new_word}】" if not new_word.startswith('【') else new_word
                    )

        # 正则替换（获取详情）
        result = apply_regex_replacements(modified_text, return_details=True)
        modified_text, regex_count, regex_stats, details = result

        # 调整详情的偏移量
        for d in details:
            d.start += global_offset
            d.end += global_offset
            para_matches.append(d)

        for name, c in regex_stats.items():
            stats[name] = stats.get(name, 0) + c

        all_matches.extend(para_matches)
        replaced_paragraphs.append(modified_text)
        global_offset += len(original_text) + 1  # +1 for paragraph separator

    return {
        "original_text": "\n".join(original_paragraphs),
        "replaced_text": "\n".join(replaced_paragraphs),
        "matches": [asdict(m) for m in all_matches],
        "stats": stats,
        "total_count": len(all_matches)
    }

# ============ API 路由 ============

@app.post("/preview")
async def preview_file(
    file: UploadFile = File(...),
    rules: str = Form("[]")
):
    """预览文档替换结果（不保存文件）"""
    try:
        rules_list = json.loads(rules) if rules else []
        rules_df = pd.DataFrame(rules_list) if rules_list else pd.DataFrame(columns=['原词', '替换词'])

        content = await file.read()
        doc_file = io.BytesIO(content)

        result = preview_document(doc_file, rules_df)

        return JSONResponse({
            "success": True,
            "filename": file.filename,
            **result
        })

    except Exception as e:
        return JSONResponse({
            "success": False,
            "error": str(e)
        }, status_code=500)

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    rules: str = Form("[]"),
    selectedMatches: str = Form("[]")  # 用户选中的匹配项ID列表
):
    """上传并处理 Word 文档"""
    try:
        # 解析规则
        import json
        rules_list = json.loads(rules) if rules else []
        rules_df = pd.DataFrame(rules_list) if rules_list else pd.DataFrame(columns=['原词', '替换词'])

        # 解析选中的匹配项ID（暂时未使用，待实现过滤逻辑）
        selected_match_ids = set(json.loads(selectedMatches) if selectedMatches else [])
        print(f"选中的匹配项ID: {selected_match_ids}")  # 调试信息

        # 读取文件
        content = await file.read()
        doc_file = io.BytesIO(content)

        # 处理文档（暂时忽略 selected_match_ids，待实现过滤逻辑）
        doc, excel_count, regex_stats = process_document(doc_file, rules_df, selected_match_ids)

        # 保存到临时文件
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        # 生成临时文件名
        temp_dir = tempfile.gettempdir()
        output_filename = f"{uuid.uuid4().hex}_已替换.docx"
        output_path = os.path.join(temp_dir, output_filename)

        with open(output_path, "wb") as f:
            f.write(output.getvalue())

        # 保存原文件名映射
        download_name = file.filename.replace('.docx', '_已替换.docx')
        temp_file_mapping[output_filename] = download_name

        total_regex = sum(regex_stats.values())

        return JSONResponse({
            "success": True,
            "filename": file.filename.replace('.docx', '_已替换.docx'),
            "temp_path": output_filename,
            "stats": {
                "excel_rules": excel_count,
                "regex_total": total_regex,
                "regex_details": regex_stats
            }
        })

    except Exception as e:
        return JSONResponse({
            "success": False,
            "error": str(e)
        }, status_code=500)

@app.get("/download/{filename}")
async def download_file(filename: str):
    """下载处理后的文件"""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, filename)

    if os.path.exists(file_path):
        # 使用映射获取正确的下载文件名
        download_name = temp_file_mapping.get(filename, filename)
        return FileResponse(
            file_path,
            filename=download_name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        return JSONResponse({"error": "文件不存在"}, status_code=404)

@app.get("/")
async def root():
    """根路由 - 返回前端页面"""
    return FileResponse("index.html")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
